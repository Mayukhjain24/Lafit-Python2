import streamlit as st
import pandas as pd
from docx import Document
import re
from fuzzywuzzy import process, fuzz
from io import BytesIO
import zipfile
import os
import tempfile

# Sanitize filenames
def sanitize_filename(name):
    return re.sub(r'[^\w\-_\.]', '_', str(name))

# Normalize text for matching
def normalize_text(text):
    return ''.join(char for char in text if char.isalnum()).lower()

# Extract placeholders from the Word document
def extract_placeholders(doc):
    placeholders = set()
    pattern = re.compile(r'\{\{(.+?)\}\}')
    def get_full_text(runs):
        return ''.join(run.text for run in runs)
    for paragraph in doc.paragraphs:
        full_text = get_full_text(paragraph.runs)
        matches = pattern.findall(full_text)
        placeholders.update(matches)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = get_full_text(paragraph.runs)
                    matches = pattern.findall(full_text)
                    placeholders.update(matches)
    for section in doc.sections:
        for header in section.header.paragraphs:
            full_text = get_full_text(header.runs)
            matches = pattern.findall(full_text)
            placeholders.update(matches)
        for footer in section.footer.paragraphs:
            full_text = get_full_text(footer.runs)
            matches = pattern.findall(full_text)
            placeholders.update(matches)
    return placeholders

# Find strings with { or } that are not placeholders
def find_invalid_braces(doc):
    invalid_strings = set()
    pattern = re.compile(r'(?<!\{)\{[^}]*\}|[^{]*\}[^{]*')
    def check_text(text):
        matches = pattern.findall(text)
        if matches:
            invalid_strings.update(matches)
    for paragraph in doc.paragraphs:
        check_text(''.join(run.text for run in paragraph.runs))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    check_text(''.join(run.text for run in paragraph.runs))
    for section in doc.sections:
        for header in section.header.paragraphs:
            check_text(''.join(run.text for run in header.runs))
        for footer in section.footer.paragraphs:
            check_text(''.join(run.text for run in footer.runs))
    return invalid_strings

# Fuzzy match placeholders to Excel columns
def fuzzy_match_placeholders(placeholders, columns, threshold=85):
    mapping = {}
    for placeholder in placeholders:
        norm_placeholder = normalize_text(placeholder)
        norm_columns = {col: normalize_text(col) for col in columns}
        best_match, score = process.extractOne(
            norm_placeholder, 
            list(norm_columns.values()), 
            scorer=fuzz.token_sort_ratio
        )
        if score >= threshold:
            for col, norm_col in norm_columns.items():
                if norm_col == best_match:
                    mapping[placeholder] = (col, score)
                    break
        else:
            mapping[placeholder] = (None, 0)
    return mapping

# Replace text in a paragraph
def replace_text_in_paragraph(paragraph, old_text, new_text):
    full_text = ''.join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return
    start_pos = full_text.find(old_text)
    end_pos = start_pos + len(old_text)
    current_pos = 0
    for run in paragraph.runs:
        run_start = current_pos
        run_end = run_start + len(run.text)
        if run_start < end_pos and run_end > start_pos:
            overlap_start = max(run_start, start_pos)
            overlap_end = min(run_end, end_pos)
            keep_before = run.text[:overlap_start - run_start] if overlap_start > run_start else ''
            keep_after = run.text[overlap_end - run_start:] if overlap_end < run_end else ''
            if run_start <= start_pos < run_end:
                run.text = keep_before + new_text + keep_after
            else:
                run.text = keep_before + keep_after
        current_pos = run_end

# Replace placeholders in the document
def replace_placeholders(doc, data, mapping):
    for placeholder, (column, _) in mapping.items():
        if column:
            placeholder_text = f"{{{{{placeholder}}}}}"
            value = str(data[column])
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, placeholder_text, value)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, placeholder_text, value)
            for section in doc.sections:
                for header in section.header.paragraphs:
                    replace_text_in_paragraph(header, placeholder_text, value)
                for footer in section.footer.paragraphs:
                    replace_text_in_paragraph(footer, placeholder_text, value)
    unreplaced = set()
    pattern = re.compile(r'\{\{(.+?)\}\}')
    def check_unreplaced(runs):
        full_text = ''.join(run.text for run in runs)
        return pattern.findall(full_text)
    for paragraph in doc.paragraphs:
        unreplaced.update(check_unreplaced(paragraph.runs))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    unreplaced.update(check_unreplaced(paragraph.runs))
    for section in doc.sections:
        for header in section.header.paragraphs:
            unreplaced.update(check_unreplaced(header.runs))
        for footer in section.footer.paragraphs:
            unreplaced.update(check_unreplaced(footer.runs))
    return doc, unreplaced

# Generate documents for each row
def generate_documents(df, template, mapping):
    documents = []
    for index, row in df.iterrows():
        temp_doc = Document(template)
        data = row.to_dict()
        temp_doc, unreplaced = replace_placeholders(temp_doc, data, mapping)
        replaced_placeholders = set(mapping.keys()) - unreplaced
        used_columns = {mapping[p][0] for p in replaced_placeholders if mapping[p][0]}
        unused_columns = set(df.columns) - used_columns
        invalid_braces = find_invalid_braces(temp_doc)
        output = BytesIO()
        temp_doc.save(output)
        output.seek(0)
        documents.append((output, replaced_placeholders, unreplaced, unused_columns, invalid_braces, row))
    return documents

# Main function
def main():
    st.title("Easy Word Document Maker")
    st.info("""
    Welcome! This app helps you create Word documents from an Excel file super easily. 
    Just upload your Excel file and a Word template with placeholders like {{Product Name}}.
    The app will fill in the details for you and let you download everything!
    """)

    # Placeholder guidance
    with st.expander("How to Set Up Your Files (Click Here!)"):
        st.write("""
        ### Super Simple Steps:
        1. **Excel File**: Use these column names (you don’t need all of them, just what you want!):
           - Product Name, Product Code, Product Category, W1, W2, W3, W4, W5, W6, W7, Product Features
           - LO(W) 1, LO(W) 2, ..., LO(W) 7, LO(N) 1, ..., LO(N) 7, LO(C) 1, ..., LO(C) 7
           - LO(TW) 1, ..., LO(TW) 7, (Lm/W) 1, ..., (Lm/W) 7, Type/No. 1, ..., Type/No. 7
           - UGR 1, ..., UGR 7, CRI, R9 Value, SDCM, LED Make, Life Hours, Temp, IP Rating
           - Optics, Beam Angles, DC 1, ..., DC 7, F1, ..., F7, Driver Brand, DR 1, ..., DR 7
           - PF 1, ..., PF 7, THD 1, ..., THD 7, Surge 1, ..., Surge 7, Class Type1, ..., Class Type7
           - IT1, ..., IT7, Automation, Input, Housing, Product Colours, Finish Type, Ral Code
           - X1, ..., X7, H1, ..., H7, Ø1, ..., Ø7, Sw.Ang1, ..., Sw.Ang7
        2. **Word Template**: Add placeholders like this: {{Product Name}}, {{Product Code}}, {{W1}}, etc.
           - Use the exact column names from your Excel file inside {{ }}.
           - Example: If your Excel has "Product Name", use {{Product Name}} in your Word file.
        3. That’s it! Upload them below and press the button!
        """)

    # File upload
    excel_file = st.file_uploader("Step 1: Upload Your Excel File", type=["xlsx"], help="Pick your Excel file with the data!")
    template_file = st.file_uploader("Step 2: Upload Your Word Template", type=["docx"], help="Pick your Word file with placeholders!")

    if excel_file and template_file:
        # Read Excel
        df = pd.read_excel(excel_file)
        columns = df.columns.tolist()
        
        # Load Word template
        doc = Document(template_file)
        placeholders = extract_placeholders(doc)
        
        # Show what we found
        st.subheader("What We Found in Your Files")
        st.write("**Placeholders in Your Word Template:**", ", ".join(placeholders) if placeholders else "None found!")
        
        # Match placeholders to columns
        mapping = fuzzy_match_placeholders(placeholders, columns)
        st.write("**How Placeholders Match Your Excel Columns:**")
        for p, (col, score) in mapping.items():
            st.write(f"- {p} → {col if col else 'No Match'} (Match Score: {score})")

        # Option to show missed columns
        show_missed_columns = st.checkbox("Show Columns Not Used in Documents", help="Check to see which Excel columns weren't used.")

        if st.button("Step 3: Make My Documents!"):
            with st.spinner("Making your documents, please wait..."):
                documents = generate_documents(df, template_file, mapping)
            
            # Create ZIP file with named documents
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                for idx, (doc_data, _, _, _, _, row) in enumerate(documents):
                    product_name = sanitize_filename(row.get('Product Name', 'unknown'))
                    product_code = sanitize_filename(row.get('Product Code', 'unknown'))
                    filename = f"{product_name}_{product_code}_row_{idx + 1}.docx"
                    zip_file.writestr(filename, doc_data.getvalue())
            zip_buffer.seek(0)
            
            # Download all as ZIP
            st.success("All done! Download your files below!")
            st.info("Documents are named as 'product_name_product_code_row_X.docx' to ensure uniqueness.")
            st.download_button(
                label="Download All Documents in a ZIP",
                data=zip_buffer,
                file_name="all_documents.zip",
                mime="application/zip",
                help="Click to get all your documents in one ZIP file!"
            )
            
            # Show concise reports
            for idx, (doc_data, replaced, unreplaced, unused_columns, invalid_braces, row) in enumerate(documents):
                product_name = row.get('Product Name', 'Not found')
                product_code = row.get('Product Code', 'Not found')
                filename = f"{sanitize_filename(product_name)}_{sanitize_filename(product_code)}_row_{idx + 1}.docx"
                st.subheader(f"Report for {filename}")
                
                # Summary
                st.write(f"**Product Name:** {product_name}")
                st.write(f"**Product Code:** {product_code}")
                st.write(f"**Replaced Placeholders:** {len(replaced)} found, {len(unreplaced)} not replaced")
                
                # Expandable detailed report
                with st.expander("See Full Details"):
                    # Hierarchical structure
                    st.write("**Product Details:**")
                    for field in ['Product Name', 'Product Code', 'Product Category', 'Product Features']:
                        if field in row:
                            st.write(f"- **{field}:** {row[field]}")
                    
                    st.write("**Wattage and Lumen Output:**")
                    for field in ['W1', 'W2', 'W3', 'W4', 'W5', 'W6', 'W7']:
                        if field in row:
                            st.write(f"- **{field}:** {row[field]}")
                    
                    for prefix in ['LO(W)', 'LO(N)', 'LO(C)', 'LO(TW)', '(Lm/W)', 'Type/No.', 'UGR', 'DC', 'F', 'DR', 'PF', 'THD', 'Surge', 'Class Type', 'IT', 'X', 'H', 'Ø', 'Sw.Ang']:
                        st.write(f"**{prefix} Values:**")
                        for i in range(1, 8):
                            col = f"{prefix} {i}" if prefix != '(Lm/W)' else f"(Lm/W) {i}"
                            if col in row:
                                st.write(f"- **{col}:** {row[col]}")
                    
                    st.write("**Additional Specifications:**")
                    for field in ['CRI', 'R9 Value', 'SDCM', 'LED Make', 'Life Hours', 'Temp', 'IP Rating', 'Optics', 'Beam Angles', 'Driver Brand', 'Automation', 'Input', 'Housing', 'Product Colours', 'Finish Type', 'Ral Code']:
                        if field in row:
                            st.write(f"- **{field}:** {row[field]}")
                    
                    # Placeholder details
                    st.write("**Placeholders Found and Replaced:**")
                    if replaced:
                        for p in sorted(replaced):
                            col, score = mapping[p]
                            if col:
                                st.write(f"- {p} → {row[col]} (from '{col}')")
                    else:
                        st.write("- None")
                    
                    st.write("**Placeholders Not Replaced:**")
                    st.write(f"- {', '.join(unreplaced) if unreplaced else 'None'}")
                    
                    if show_missed_columns:
                        st.write("**Excel Columns Not Used:**")
                        st.write(f"- {', '.join(sorted(unused_columns)) if unused_columns else 'None'}")
                    
                    st.write("**Strings with { or } Found:**")
                    st.write(f"- {', '.join(invalid_braces) if invalid_braces else 'None'}")
                
                # Download individual document
                st.download_button(
                    label=f"Download {filename}",
                    data=doc_data,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help=f"Click to download {filename}"
                )

if __name__ == "__main__":
    main()